"""DOCX text extraction via python-docx.

python-docx reads the document as a sequence of paragraphs and tables.
We flatten both into plain text, keeping paragraph order and rendering
tables as TSV (tab-separated) rows so retrieval can still match cell
content.

DOCX has no real page concept (pagination is a Word render-time
decision), so we emit a single ParsedPage. If we ever need section/
heading anchoring, that comes from paragraph.style.name.
"""

from __future__ import annotations

from io import BytesIO

import docx  # python-docx exposes itself as `docx`

from rag_service.parsers.types import ParsedDocument, ParsedPage


def _render_table(table: docx.table.Table) -> str:
    """Serialize a Word table as TSV: rows separated by \\n, cells by \\t."""
    rows: list[str] = []
    for row in table.rows:
        cells = [cell.text.replace("\t", " ").replace("\n", " ") for cell in row.cells]
        rows.append("\t".join(cells))
    return "\n".join(rows)


def parse_docx(content: bytes) -> ParsedDocument:
    """Extract paragraphs + tables from a .docx byte stream."""
    document = docx.Document(BytesIO(content))

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        rendered = _render_table(table).strip()
        if rendered:
            blocks.append(rendered)

    full_text = "\n\n".join(blocks)
    return ParsedDocument(
        full_text=full_text,
        pages=[ParsedPage(page_number=1, text=full_text)],
        parser="python-docx",
        extra={
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
    )
